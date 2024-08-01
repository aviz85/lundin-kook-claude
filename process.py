import os
import json
import logging
import argparse
import sys
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import anthropic
from datetime import datetime

# Load environment variables
load_dotenv()

# Choose model: 'claude-3-opus-20240229' or 'claude-3-5-sonnet-20240620'
MODEL = 'claude-3-5-sonnet-20240620'

# Set up Anthropic client
client = anthropic.Anthropic()

# Set up logging
def setup_logging():
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)

    # Create handlers
    c_handler = logging.StreamHandler(sys.stdout)
    f_handler = logging.FileHandler('api_usage.log')
    c_handler.setLevel(logging.DEBUG)
    f_handler.setLevel(logging.INFO)

    # Create formatters and add it to handlers
    log_format = '%(asctime)s - %(levelname)s - %(message)s'
    c_format = logging.Formatter(log_format)
    f_format = logging.Formatter(log_format)
    c_handler.setFormatter(c_format)
    f_handler.setFormatter(f_format)

    # Add handlers to the logger
    logger.addHandler(c_handler)
    logger.addHandler(f_handler)

    return logger

logger = setup_logging()

# Initialize usage counters
total_input_tokens = 0
total_output_tokens = 0
model_usage = {}
errors = []

def read_file(filename):
    logger.debug(f"Reading file: {filename}")
    with open(filename, 'r', encoding='utf-8') as file:
        content = file.read()
    logger.debug(f"File {filename} read successfully. Content length: {len(content)}")
    return content

def write_file(filename, content):
    logger.debug(f"Writing to file: {filename}")
    with open(filename, 'w', encoding='utf-8') as file:
        json.dump(content, file, ensure_ascii=False, indent=2)
    logger.debug(f"Content written to {filename} successfully")

def call_claude_api(system_prompt, user_message):
    global total_input_tokens, total_output_tokens, model_usage
    
    logger.debug(f"Preparing API call with system prompt length: {len(system_prompt)} and user message length: {len(user_message)}")
    
    try:
        logger.debug("Sending request to Claude API")
        response = client.messages.create(
            model=MODEL,
            max_tokens=4096,
            messages=[
                {"role": "user", "content": user_message}
            ],
            system=system_prompt
        )
        
        # Log usage
        input_tokens = response.usage.input_tokens
        output_tokens = response.usage.output_tokens
        total_input_tokens += input_tokens
        total_output_tokens += output_tokens
        
        # Update model-specific usage
        if MODEL not in model_usage:
            model_usage[MODEL] = {'input_tokens': 0, 'output_tokens': 0}
        model_usage[MODEL]['input_tokens'] += input_tokens
        model_usage[MODEL]['output_tokens'] += output_tokens
        
        logger.info(f"API call - Model: {MODEL}, Input tokens: {input_tokens}, Output tokens: {output_tokens}")
        
        # Parse the response content
        content = json.loads(response.content[0].text)
        
        # Add usage data to the content
        content['usage'] = {
            'input_tokens': input_tokens,
            'output_tokens': output_tokens
        }
        
        return content
    except anthropic.APIError as e:
        error_msg = f"Error in API call: {str(e)}"
        errors.append(error_msg)
        logger.error(error_msg)
        return None

def set_rtl(paragraph):
    logger.debug("Setting RTL for paragraph")
    pPr = paragraph._p.get_or_add_pPr()
    biDi = OxmlElement('w:bidi')
    pPr.insert_element_before(biDi, 'w:jc')

def compile_to_docx(results_dir, output_file):
    logger.info(f"Compiling results from {results_dir} to {output_file}")
    doc = Document()
    
    # Set RTL direction for the entire document
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    for filename in sorted(os.listdir(results_dir)):
        if filename.endswith('.json'):
            logger.debug(f"Processing file: {filename}")
            try:
                with open(os.path.join(results_dir, filename), 'r', encoding='utf-8') as file:
                    data = json.load(file)
                
                # Add letter
                para = doc.add_paragraph()
                set_rtl(para)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run(data.get('letter', ''))
                run.bold = True
                doc.add_paragraph()
                
                # Add original text
                para = doc.add_paragraph()
                set_rtl(para)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                para.add_run(data['original_text'])
                doc.add_paragraph()
                
                # Add difficult words explanations
                difficult_words = data['difficult_words']
                if difficult_words:
                    para = doc.add_paragraph()
                    set_rtl(para)
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    explanations = '; '.join([f"{item['word']} – {item['explanation']}" for item in difficult_words])
                    para.add_run(explanations)
                    doc.add_paragraph()
                
                # Add detailed interpretation
                detailed_interpretation = data['detailed_interpretation']
                para = doc.add_paragraph()
                set_rtl(para)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                for part in detailed_interpretation:
                    run = para.add_run(part['quote'])
                    run.bold = True
                    para.add_run(f" - {part['explanation']} ")
                
                # Add extra paragraph for spacing between sections
                doc.add_paragraph()
                logger.debug(f"Finished processing file: {filename}")
            except Exception as e:
                error_msg = f"Error processing file {filename}: {str(e)}"
                errors.append(error_msg)
                logger.error(error_msg)
    
    doc.save(output_file)
    logger.info(f"Compiled document saved as {output_file}")

def main():
    parser = argparse.ArgumentParser(description="Process text files and compile into DOCX.")
    parser.add_argument("--skip-processing", action="store_true", help="Skip processing and only compile existing JSON files")
    args = parser.parse_args()

    logger.info("Script execution started")

    if not args.skip_processing:
        logger.info("Processing mode: Full processing")
        # Read prompt and examples
        prompt = read_file('prompt.txt')
        examples = read_file('examples.txt')
        
        # Example of correct interpretation structure
        example_structure = {
            "letter": "א",
            "original_text": "התכונה של יראת שמים, מצד עצמה, לית לה מגרמה כלום, ואי אפשר לה להיות מתחשבת בין הכשרונות ומעלות הנפש של האדם.",
            "difficult_words": [
                {"word": "לית לה מגרמה כלום", "explanation": "אין לה מעצמה כלום"},
                {"word": "מתחשבת", "explanation": "נחשבת, נספרת"}
            ],
            "detailed_interpretation": [
                {
                    "quote": "התכונה של יראת שמים, מצד עצמה, לית לה מגרמה כלום",
                    "explanation": "השאיפה הדתית (\"יראת שמים\") איננה תוכן העומד בפני עצמו. הרב קוק מסביר כי יראת שמים, כשלעצמה, אינה בעלת ערך עצמאי."
                },
                {
                    "quote": "ואי אפשר לה להיות מתחשבת בין הכשרונות ומעלות הנפש של האדם",
                    "explanation": "יראת שמים איננה נספרת בין שאר כוחות הנפש. היא אינה יכולה להיחשב כאחת מהתכונות או היכולות של האדם."
                }
            ]
        }
        
        # Combine prompt, examples, and JSON structure example
        system_prompt = f"""{prompt}

    פסקאות לדוגמא:

    {examples}

    Please provide your interpretation in JSON format. Here's an example of the correct structure:

    {json.dumps(example_structure, ensure_ascii=False, indent=2)}

    Make sure to follow this structure in your response, using JSON mode."""
        
        # Log the system prompt
        logger.debug("System Prompt prepared")
        print("System Prompt:")
        print("=" * 50)
        print(system_prompt[:500] + "..." if len(system_prompt) > 500 else system_prompt)
        print("=" * 50)
        
        # Create results directory if it doesn't exist
        os.makedirs('results', exist_ok=True)
        logger.debug("Results directory created/confirmed")
        
        # Process all txt files in the sources directory
        for filename in os.listdir('sources'):
            if filename.endswith('.txt'):
                logger.info(f"Processing file: {filename}")
                print(f"\nProcessing {filename}...")
                
                # Read the paragraph
                paragraph = read_file(os.path.join('sources', filename))
                
                # Log the current paragraph
                logger.debug(f"Paragraph from {filename}: {paragraph[:100]}...")
                print("Current Paragraph:")
                print("-" * 50)
                print(paragraph[:500] + "..." if len(paragraph) > 500 else paragraph)
                print("-" * 50)
                
                # Call Claude API
                print("Calling Claude API...")
                response = call_claude_api(system_prompt, paragraph)
                
                if response:
                    # Generate timestamp
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    
                    # Determine output filename with timestamp
                    output_filename = f"results/{os.path.splitext(filename)[0]}_{MODEL}_{timestamp}.json"
                    
                    # Write the response to a file
                    write_file(output_filename, response)
                    
                    logger.info(f"Result saved to {output_filename}")
                    print(f"Result saved to {output_filename}")
                    
                    # Log a snippet of the response
                    logger.debug(f"Response snippet: {json.dumps(response, ensure_ascii=False)[:100]}...")
                    print("Response snippet:")
                    print("~" * 50)
                    print(json.dumps(response, ensure_ascii=False, indent=2)[:500] + "..." if len(json.dumps(response, ensure_ascii=False)) > 500 else json.dumps(response, ensure_ascii=False, indent=2))
                    print("~" * 50)
                else:
                    logger.warning(f"Failed to process {filename}")
                    print(f"Failed to process {filename}")
    else:
        logger.info("Processing mode: Skip processing, compile only")
    
    # Compile all JSON files into a single DOCX
    compile_to_docx('results', 'compiled_interpretations.docx')
    
    # Log total usage
    logger.info("Total usage:")
    for model, usage in model_usage.items():
        logger.info(f"Model: {model} - Input tokens: {usage['input_tokens']}, Output tokens: {usage['output_tokens']}")
    logger.info(f"Overall - Input tokens: {total_input_tokens}, Output tokens: {total_output_tokens}")
    
    print("Total usage:")
    for model, usage in model_usage.items():
        print(f"Model: {model} - Input tokens: {usage['input_tokens']}, Output tokens: {usage['output_tokens']}")
    print(f"Overall - Input tokens: {total_input_tokens}, Output tokens: {total_output_tokens}")
    
    if errors:
        logger.warning("Errors encountered during processing")
        print("\nErrors encountered during processing:")
        for error in errors:
            print(error)

    logger.info("Script execution completed")
    print("\nScript execution completed. Check the output above for detailed logs.")

if __name__ == "__main__":
    main()